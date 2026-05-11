"""
process_figures.py — KAKENHI Pipeline Figure Extraction & Embedding Tool v2.1.0

End-to-end automation:
  1. extract-pages: PDF → per-page PNGs (audit trail)
  2. crop-rename:   Select + crop + rename key figures
  3. embed-md:      Inject figure references into cf-19 .md files
  4. embed-docx:    Inject figures into cf-19 .docx files via python-docx
  5. auto:          Full pipeline (extract → crop → embed) for a grant folder

Dependencies: Pillow, python-docx, poppler-utils (pdftoppm)
"""

from PIL import Image, ImageChops
import os, sys, argparse, subprocess, shutil, glob, re, json
from pathlib import Path


def trim(im):
    """Trim whitespace/solid background from image borders."""
    bg = Image.new(im.mode, im.size, im.getpixel((0, 0)))
    diff = ImageChops.difference(im, bg)
    diff = ImageChops.add(diff, diff, 2.0, -100)
    bbox = diff.getbbox()
    return im.crop(bbox) if bbox else im


def generate_captions_with_gemini(image_path, paper_name):
    """Use Gemini 2.5 Pro Vision to extract the figure caption and generate JP/EN summaries."""
    try:
        from google import genai
        from google.genai import types
    except ImportError:
        print("  WARN: google.genai not installed. Falling back to hardcoded captions.")
        return None, None

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        env_file = os.path.expanduser("~/.gemini/.env")
        if os.path.exists(env_file):
            with open(env_file, "r") as f:
                for line in f:
                    if line.startswith("GEMINI_API_KEY=") or line.startswith("GOOGLE_API_KEY="):
                        api_key = line.strip().split("=", 1)[1]
                        break

    if not api_key:
        print("  WARN: No GEMINI_API_KEY found. Falling back to hardcoded captions.")
        return None, None

    print("  [Gemini] Analyzing figure to generate captions...")
    try:
        client = genai.Client(api_key=api_key)
        im = Image.open(image_path)
        
        prompt = f"""
You are a scientific academic assistant. 
Analyze the provided page/figure from a scientific paper (Source: {paper_name}).
1. Identify the main figure on this page (e.g., "Figure 3").
2. Extract the original English caption.
3. Write a short, single-sentence Japanese caption suitable for a KAKENHI report (e.g., "図3. [Brief description] ([First Author] et al., [Journal] [Year]より)"). If journal/year is unknown, just use the paper filename.
4. Write a short, single-sentence English caption.

Respond ONLY with a valid JSON object:
{{
  "caption_jp": "図X. ...",
  "caption_en": "Figure X. ..."
}}
"""
        response = client.models.generate_content(
            model='gemini-2.5-pro',
            contents=[prompt, im],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
            ),
        )
        data = json.loads(response.text)
        return data.get("caption_jp"), data.get("caption_en")
    except Exception as e:
        print(f"  WARN: Gemini API call failed: {e}")
        return None, None


def extract_pages(pdf_path, output_dir, dpi=200):
    """Extract all pages from a PDF as PNG images using pdftoppm."""
    os.makedirs(output_dir, exist_ok=True)
    base = Path(pdf_path).stem[:40]  # Truncate long names
    safe = re.sub(r'[^\w\-]', '_', base)
    prefix = os.path.join(output_dir, safe)
    cmd = ["pdftoppm", "-png", "-r", str(dpi), pdf_path, prefix]
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)
        extracted = sorted([
            f for f in os.listdir(output_dir)
            if f.endswith('.png') and f.startswith(safe)
        ])
        print(f"  Extracted {len(extracted)} pages from {Path(pdf_path).name}")
        return extracted
    except FileNotFoundError:
        print("ERROR: pdftoppm not found. Install: sudo apt-get install poppler-utils")
        sys.exit(1)
    except subprocess.CalledProcessError as e:
        print(f"ERROR extracting: {e.stderr}")
        return []


def crop_and_rename(input_dir, mapping_str, output_dir=None):
    """Crop whitespace and rename figure files."""
    if output_dir is None:
        output_dir = str(Path(input_dir).parent)
    os.makedirs(output_dir, exist_ok=True)
    pairs = [pair.split(':') for pair in mapping_str.split(',')]
    results = []
    for src_name, dst_name in pairs:
        src_path = os.path.join(input_dir, src_name.strip())
        dst_path = os.path.join(output_dir, dst_name.strip())
        if os.path.exists(src_path):
            im = Image.open(src_path)
            im = trim(im)
            im.save(dst_path)
            print(f"  Cropped: {src_name.strip()} -> {dst_name.strip()}")
            results.append(dst_path)
        else:
            print(f"  WARN: Not found -> {src_path}")
    return results


def find_first_author_paper(paper_dir, pi_name_variants):
    """Find the PI's most recent first-author paper PDF (newest first)."""
    pdfs = sorted(glob.glob(os.path.join(paper_dir, "*.pdf")))
    pdfs = [p for p in pdfs if not Path(p).name.startswith("._")]
    # Sort by year prefix descending (newest first): "2025-..." before "2021-..."
    pdfs_desc = sorted(pdfs, key=lambda p: Path(p).name, reverse=True)
    # Priority 1: Most recent paper with PI surname in filename
    for pdf in pdfs_desc:
        name = Path(pdf).name.lower()
        for variant in pi_name_variants:
            if variant.lower() in name:
                return pdf
    # Priority 2: Most recent paper overall
    return pdfs_desc[0] if pdfs_desc else None


def select_best_figure_page(audit_dir, total_pages):
    """Heuristic: select the page most likely to contain a key figure.
    Scientific papers typically have main figures on pages 3-7."""
    pngs = sorted(glob.glob(os.path.join(audit_dir, "*.png")))
    if not pngs:
        return None, None
    # For short papers (<=5 pages), pick page 3 or 4
    # For longer papers, pick page 4-5
    if total_pages <= 3:
        idx = min(1, len(pngs) - 1)  # page 2
    elif total_pages <= 6:
        idx = min(2, len(pngs) - 1)  # page 3
    else:
        idx = min(3, len(pngs) - 1)  # page 4
    return pngs[idx], Path(pngs[idx]).name


def embed_figure_in_md(md_path, figure_rel_path, caption, section_marker="４．研究成果"):
    """Insert figure markdown reference into cf-19 .md file after achievements section header."""
    if not os.path.exists(md_path):
        print(f"  WARN: MD file not found: {md_path}")
        return False
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    fig_md = f"\n\n![{caption}]({figure_rel_path})\n\n*{caption}*\n\n"
    # For English version, try different markers
    markers = [section_marker, "4. Research Achievements", "**4. Research Achievements**",
               "４．研究成果", "(1)", "（1）"]
    inserted = False
    for marker in markers:
        if marker in content:
            # Insert figure after the first finding paragraph
            pos = content.find(marker)
            # Find the end of the first paragraph after the marker
            next_para = content.find("\n\n", pos + len(marker))
            if next_para > 0:
                next_para2 = content.find("\n\n", next_para + 2)
                insert_pos = next_para2 if next_para2 > 0 else next_para
                content = content[:insert_pos] + fig_md + content[insert_pos:]
                inserted = True
                break
    if not inserted:
        # Fallback: append before references
        content += fig_md
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"  Embedded figure in {Path(md_path).name}")
    return True


def embed_figure_in_docx(docx_path, figure_path, caption):
    """Insert figure into cf-19 .docx file using python-docx."""
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  ERROR: python-docx not installed. Run: pip install python-docx")
        return False
    if not os.path.exists(docx_path) or not os.path.exists(figure_path):
        print(f"  WARN: File not found: {docx_path} or {figure_path}")
        return False
    doc = Document(docx_path)
    # Find the paragraph containing "４．研究成果" or "(1)"
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "４．研究成果" in text or "4. Research" in text or "研究成果" in text:
            target_idx = i
            break
    if target_idx is None:
        # Fallback: insert near the end
        target_idx = max(0, len(doc.paragraphs) - 3)
    # Find the paragraph after the first finding
    insert_idx = None
    for i, para in enumerate(doc.paragraphs):
        if i > target_idx:
            text = para.text.strip()
            if text.startswith('（2）') or text.startswith('(2)'):
                insert_idx = i
                break
    if insert_idx is None:
        insert_idx = min(target_idx + 3, len(doc.paragraphs) - 1)

    # Use add_picture at document level (appended), then move it
    doc.add_picture(figure_path, width=Cm(14))
    pic_para = doc.paragraphs[-1]._element
    pic_para.getparent().remove(pic_para)

    # Use add_paragraph for caption (appended), then move it
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    cap_element = cap_para._element
    cap_element.getparent().remove(cap_element)

    # Insert at the correct position (above the (2) paragraph)
    # The order we insert is: picture then caption
    if insert_idx < len(doc.paragraphs):
        doc.paragraphs[insert_idx]._element.addprevious(pic_para)
        doc.paragraphs[insert_idx]._element.addprevious(cap_element)
    
    doc.save(docx_path)
    print(f"  Embedded figure in {Path(docx_path).name}")
    return True


def auto_pipeline(grant_folder, pi_name_variants=None, grant_id=None, report_year=None):
    """Full automated pipeline: extract → select → crop → embed."""
    if pi_name_variants is None:
        pi_name_variants = ["Weng"]
    paper_dir = os.path.join(grant_folder, "References", "paper")
    figures_dir = os.path.join(grant_folder, "References", "figures")
    audit_dir = os.path.join(figures_dir, "audit")
    reports_dir = os.path.join(grant_folder, "Reports")
    os.makedirs(audit_dir, exist_ok=True)
    print(f"=== KAKENHI Figure Pipeline v2.1 ===")
    print(f"Grant folder: {grant_folder}")

    # 1. Find best paper
    pdf = find_first_author_paper(paper_dir, pi_name_variants)
    if not pdf:
        print("ERROR: No PDFs found in References/paper/")
        return False
    print(f"\n[1/5] Selected paper: {Path(pdf).name}")

    # 2. Extract pages
    print(f"\n[2/5] Extracting pages...")
    pages = extract_pages(pdf, audit_dir, dpi=200)
    if not pages:
        print("ERROR: No pages extracted")
        return False

    # 3. Select best figure page
    print(f"\n[3/5] Selecting key figure...")
    best_png, best_name = select_best_figure_page(audit_dir, len(pages))
    if not best_png:
        print("ERROR: No figure selected")
        return False
    print(f"  Selected: {best_name} (page {pages.index(best_name)+1} of {len(pages)})")

    # 4. Crop and save
    print(f"\n[4/5] Cropping and saving...")
    figure_filename = "Figure1_OGlcNAc_Organelle.png"
    mapping = f"{best_name}:{figure_filename}"
    results = crop_and_rename(audit_dir, mapping, figures_dir)

    # 5. Embed in reports
    print(f"\n[5/5] Generating captions and embedding in reports...")
    figure_abs = os.path.join(figures_dir, figure_filename)
    
    cap_jp, cap_en = generate_captions_with_gemini(figure_abs, Path(pdf).name)
    
    caption_jp = cap_jp if cap_jp else "図1. 抽出された代表的画像（Weng Y et al.より）"
    caption_en = cap_en if cap_en else "Figure 1. Representative extracted image (from Weng Y et al.)"
    
    print(f"  Generated JP: {caption_jp}")
    print(f"  Generated EN: {caption_en}")
    
    figure_rel = f"../References/figures/{figure_filename}"

    # Determine file prefix
    prefix = f"{grant_id}_{report_year}" if grant_id and report_year else ""
    if not prefix:
        # Auto-detect from existing files
        for f in os.listdir(reports_dir):
            m = re.match(r'(\w+_\d{4})_F-19-1', f)
            if m:
                prefix = m.group(1)
                break

    if prefix:
        jp_md = os.path.join(reports_dir, f"{prefix}_F-19-1-cf-19_draft.md")
        en_md = os.path.join(reports_dir, f"{prefix}_F-19-1-cf-19_draft_en.md")
        jp_docx = os.path.join(reports_dir, f"{prefix}_F-19-1-cf-19_draft.docx")
        en_docx = os.path.join(reports_dir, f"{prefix}_F-19-1-cf-19_draft_en.docx")

        embed_figure_in_md(jp_md, figure_rel, caption_jp)
        embed_figure_in_md(en_md, figure_rel, caption_en, "4. Research Achievements")
        embed_figure_in_docx(jp_docx, figure_abs, caption_jp)
        embed_figure_in_docx(en_docx, figure_abs, caption_en)

    # Write audit log
    log_path = os.path.join(audit_dir, "audit_log.md")
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(f"# Figure Extraction Audit Log\n\n")
        f.write(f"- **Date**: {__import__('datetime').datetime.now().isoformat()}\n")
        f.write(f"- **Source PDF**: {Path(pdf).name}\n")
        f.write(f"- **Pages extracted**: {len(pages)}\n")
        f.write(f"- **Selected page**: {best_name} (page {pages.index(best_name)+1})\n")
        f.write(f"- **Output figure**: {figure_filename}\n")
        f.write(f"- **Caption (JP)**: {caption_jp}\n")
        f.write(f"- **Caption (EN)**: {caption_en}\n")
        f.write(f"- **Embedded in**: cf-19 JP/EN .md and .docx\n")
    print(f"\n✅ Audit log written: {log_path}")
    print(f"✅ Figure pipeline complete!")
    return True


def main():
    parser = argparse.ArgumentParser(description="KAKENHI Pipeline Figure Tool v2.1")
    subparsers = parser.add_subparsers(dest="command")

    ep = subparsers.add_parser("extract-pages", help="PDF → PNGs")
    ep.add_argument("--pdf", required=True)
    ep.add_argument("--output", required=True)
    ep.add_argument("--dpi", type=int, default=200)

    cr = subparsers.add_parser("crop-rename", help="Crop + rename")
    cr.add_argument("--input", required=True)
    cr.add_argument("--output", default=None)
    cr.add_argument("--mapping", required=True)

    ap = subparsers.add_parser("auto", help="Full automated pipeline")
    ap.add_argument("--grant-folder", required=True)
    ap.add_argument("--pi-names", default="Weng", help="Comma-separated PI name variants")
    ap.add_argument("--grant-id", default=None)
    ap.add_argument("--report-year", default=None)

    args = parser.parse_args()
    if args.command == "extract-pages":
        extract_pages(args.pdf, args.output, args.dpi)
    elif args.command == "crop-rename":
        crop_and_rename(args.input, args.mapping, args.output)
    elif args.command == "auto":
        pi_names = [n.strip() for n in args.pi_names.split(',')]
        auto_pipeline(args.grant_folder, pi_names, args.grant_id, args.report_year)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
